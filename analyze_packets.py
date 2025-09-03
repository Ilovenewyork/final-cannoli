import docx
import os

def analyze_docx_structure(docx_path):
    print(f"\nAnalyzing structure of: {os.path.basename(docx_path)}")
    print("=" * 60)
    
    try:
        doc = docx.Document(docx_path)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        # Print first 20 paragraphs with their styles and first 50 chars of text
        print("\nFirst 20 paragraphs:")
        print("-" * 60)
        for i, para in enumerate(doc.paragraphs[:20], 1):
            text = para.text.strip()
            if not text:
                continue
                
            style = para.style.name if para.style else "No Style"
            print(f"{i:3d}. [{style:15}] {text[:100]}" + ("..." if len(text) > 100 else ""))
        
        # Look for common question patterns
        print("\nPotential question indicators:")
        print("-" * 60)
        question_patterns = ["tossup", "bonus", "answer:", "<", "[", "for 10 points"]
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if any(pattern in text for pattern in question_patterns):
                print(f"Para {i+1}: {text[:150]}" + ("..." if len(text) > 150 else ""))
    
    except Exception as e:
        print(f"Error analyzing {docx_path}: {str(e)}")

if __name__ == "__main__":
    # Paths to the packet files
    farsi_packet = r"controllers\sample_packets\FARSI Packet 1.docx"
    scottie_packet = r"controllers\sample_packets\Blended Round 12 - 2021 Scottie.docx"
    
    # Analyze both packets
    analyze_docx_structure(farsi_packet)
    analyze_docx_structure(scottie_packet)
